from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_checklist():
    document = Document()

    # Title
    title = document.add_heading('Independent Living Skills Checklist', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Client Info
    p = document.add_paragraph()
    runner = p.add_run("Client's full name: __________________________________________    Date submitted: __________________")
    runner.font.size = Pt(11)
    
    document.add_paragraph() # Spacer

    # Instructions
    p = document.add_paragraph()
    runner_bold = p.add_run("Instructions: ")
    runner_bold.bold = True
    runner_text = p.add_run("Below is a long list of independent living skills divided into several categories. Please rate each item based on what you can do fine without support, what you need to practice, what you plan on starting, what you need support with, and what doesn't apply to you.")
    
    document.add_paragraph() # Spacer

    # Table
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["General Life Skills", "Can Do\nAlready", "Needs\nMore\nPractice", "Plan to\nStart", "Ongoing\nSupport\nNeeded", "N/A"]
    hdr_cells = table.rows[0].cells
    
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header_text)
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add gray shading to header
        shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Items
    items = [
        "I brush my teeth daily",
        "I shower daily with shampoo, conditioner, and soap",
        "I wash my hair",
        "I use deodorant daily",
        "I brush/comb my hair",
        "I shave as needed",
        "I can dress myself",
        "I choose the clothes I wear"
    ]

    checkbox = "☐" # Unicode empty ballot box

    for item in items:
        row_cells = table.add_row().cells
        # First column: The skill text
        row_cells[0].text = item
        
        # Other columns: Checkboxes
        for i in range(1, 6):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(checkbox)
            run.font.size = Pt(16) # Make checkbox slightly larger

    # Set column widths (approximate)
    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        for i in range(1, 6):
            row.cells[i].width = Inches(0.8)

    output_path = "Independent_Living_Skills_Checklist.docx"
    document.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_checklist()
