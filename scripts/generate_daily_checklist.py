import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION

def create_daily_checklist(json_path):
    # Load tenants data
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            tenants = json.load(f)
    except FileNotFoundError:
        print(f"Error: Could not find {json_path}")
        return

    # Sort tenants by room code just in case
    tenants.sort(key=lambda x: x.get('roomCode', ''))

    document = Document()

    # --- COVER PAGE / SUMMARY ---
    title = document.add_heading('Daily Living Skills Check - Summary Sheet', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = document.add_paragraph()
    runner = p.add_run("Date: __________________    Staff Member: __________________")
    runner.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary Table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    headers = ["Room", "Resident Name", "Done", "Notes"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add shading
        shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    for tenant in tenants:
        row_cells = table.add_row().cells
        row_cells[0].text = tenant.get('roomCode', '')
        row_cells[1].text = tenant.get('name', '')
        row_cells[2].text = "☐" # Checkbox
        row_cells[3].text = ""

    # Set column widths for summary
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(2.5)

    # --- INDIVIDUAL CHECKLISTS ---
    
    # Items for the checklist
    items = [
        "I brush my teeth daily",
        "I shower daily with shampoo, conditioner, and soap",
        "I wash my hair",
        "I use deodorant daily",
        "I brush/comb my hair",
        "I shave as needed",
        "I can dress myself",
        "I choose the clothes I wear"
    ]

    for tenant in tenants:
        document.add_section(WD_SECTION.NEW_PAGE)
        
        # Header
        title = document.add_heading('Independent Living Skills Checklist', 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Client Info - Pre-filled
        p = document.add_paragraph()
        client_info = f"Client: {tenant.get('name', '')} ({tenant.get('roomCode', '')})"
        runner = p.add_run(f"Client's full name: {client_info:<40}    Date submitted: __________________")
        runner.font.size = Pt(11)
        runner.bold = True
        
        document.add_paragraph() # Spacer

        # Instructions
        p = document.add_paragraph()
        runner_bold = p.add_run("Instructions: ")
        runner_bold.bold = True
        p.add_run("Below is a long list of independent living skills divided into several categories. Please rate each item based on what you can do fine without support, what you need to practice, what you plan on starting, what you need support with, and what doesn't apply to you.")
        
        document.add_paragraph() # Spacer

        # Table
        checklist_table = document.add_table(rows=1, cols=6)
        checklist_table.style = 'Table Grid'
        
        # Headers
        headers = ["General Life Skills", "Can Do\nAlready", "Needs\nMore\nPractice", "Plan to\nStart", "Ongoing\nSupport\nNeeded", "N/A"]
        hdr_cells = checklist_table.rows[0].cells
        
        for i, header_text in enumerate(headers):
            cell = hdr_cells[i]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header_text)
            run.bold = True
            run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        checkbox = "☐"

        for item in items:
            row_cells = checklist_table.add_row().cells
            row_cells[0].text = item
            for i in range(1, 6):
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(checkbox)
                run.font.size = Pt(14)

        # Set widths
        for row in checklist_table.rows:
            row.cells[0].width = Inches(3.0)
            for i in range(1, 6):
                row.cells[i].width = Inches(0.8)

    output_path = "Daily_Living_Skills_Batch.docx"
    document.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    # Assuming the script is run from the project root
    json_path = os.path.join("app", "tenants-data.json")
    create_daily_checklist(json_path)
