from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION

def create_landscape_tracking_sheet():
    document = Document()
    
    # Set to Landscape
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.69)  # A4 Landscape
    section.page_height = Inches(8.27)
    
    # Margins (Narrow to fit more)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Title
    title = document.add_heading('Registro Diário de Habilidades de Vida (Daily Living Skills Log)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Info Line
    p = document.add_paragraph()
    runner = p.add_run("Data: __________________    Funcionário (Staff): __________________    Turno: __________________")
    runner.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph() # Spacer

    # LEGEND / KEY
    legend_table = document.add_table(rows=1, cols=1)
    legend_table.style = 'Table Grid'
    legend_cell = legend_table.rows[0].cells[0]
    p = legend_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    runner = p.add_run("LEGENDA DE AVALIAÇÃO (RATING KEY):\n")
    runner.bold = True
    runner.font.size = Pt(10)
    
    key_text = "1 = Faz Bem (Can Do)   |   2 = Precisa Prática (Needs Practice)   |   3 = Planeja Iniciar (Plan to Start)   |   4 = Precisa Suporte (Support Needed)   |   - = N/A"
    runner = p.add_run(key_text)
    runner.font.size = Pt(10)
    
    # Add light gray background to legend
    shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
    legend_cell._tc.get_or_add_tcPr().append(shading_elm)

    document.add_paragraph() # Spacer

    # MAIN TRACKING TABLE
    # 1 Column for Skill + 10 Columns for Rooms = 11 Columns
    table = document.add_table(rows=1, cols=11)
    table.style = 'Table Grid'
    table.autofit = False 
    
    # --- HEADER ROW ---
    hdr_cells = table.rows[0].cells
    
    # Skill Column Header
    hdr_cells[0].text = "HABILIDADES / SKILLS"
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Room Columns Headers
    for i in range(1, 11):
        hdr_cells[i].text = f"Quarto\n{i}"
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].bold = True

    # Header Shading (Darker Gray)
    for cell in hdr_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- SKILL ITEMS ---
    items = [
        "Escovar os dentes (Brush teeth)",
        "Banho completo (Shower w/ soap)",
        "Lavar o cabelo (Wash hair)",
        "Usar desodorante (Use deodorant)",
        "Pentear cabelo (Brush/comb hair)",
        "Barbear-se (Shave as needed)",
        "Vestir-se sozinho (Dress myself)",
        "Escolher roupas (Choose clothes)",
        "Higiene das mãos (Hand hygiene)",
        "Cortar unhas (Trim nails)" 
    ]
    # Added a few extra common ones to fill the page nicely, user can edit

    for item in items:
        row_cells = table.add_row().cells
        
        # Skill Name
        row_cells[0].text = item
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Empty cells for 10 rooms
        for i in range(1, 11):
            row_cells[i].text = ""

    # Set Column Widths
    # Total width approx 10.5 inches usable
    # Skill col: 2.5 inches
    # Room cols: 0.8 inches each (0.8 * 10 = 8.0)
    # Total = 10.5
    
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        for i in range(1, 11):
            row.cells[i].width = Inches(0.8)

    document.add_paragraph() # Spacer

    # NOTES SECTION
    notes_title = document.add_heading('Observações / Notes:', 3)
    
    # Create a simple box for notes
    notes_table = document.add_table(rows=1, cols=1)
    notes_table.style = 'Table Grid'
    notes_cell = notes_table.rows[0].cells[0]
    notes_cell.text = "\n\n\n\n" # Empty lines for writing space

    output_path = "Planilha_Diaria_10_Quartos.docx"
    document.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_landscape_tracking_sheet()
